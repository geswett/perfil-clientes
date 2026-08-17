"""Genera el documento Word (.docx) del Perfil de Cargo a partir de los datos
estructurados (ver schema.py), replicando el estilo del formato de Puelche
Human Consulting: encabezado con logo, títulos en gris con numeración romana,
y tablas de dos columnas con la columna de etiqueta sombreada en verde claro.
"""

import io
import os
import re
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from schema import REQUISITOS_FILAS

HEADER_GRAY = RGBColor(0x80, 0x80, 0x80)
LABEL_GREEN = "D6E3BC"
WHITE = "FFFFFF"
FONT_NAME = "Calibri"
FONT_SIZE = Pt(10)

LOGO_PATH = os.path.join(os.path.dirname(__file__), "assets", "puelche_logo.jpeg")


# --------------------------------------------------------------------------
# Helpers de formato de bajo nivel
# --------------------------------------------------------------------------

def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _style_run(run, bold=False, color=None, size=FONT_SIZE):
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _set_cell_text(cell, text, bold=False, shading=None, align=None):
    _shade_cell(cell, shading if shading else WHITE)
    cell.text = ""
    lines = str(text).split("\n") if text else [""]
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(lines[0])
    _style_run(run, bold=bold)
    for line in lines[1:]:
        p2 = cell.add_paragraph()
        r2 = p2.add_run(line)
        _style_run(r2, bold=bold)


def _set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "BFBFBF")
        borders.append(el)
    tblPr.append(borders)


def _add_heading(doc, numeral, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{numeral}. {title}")
    _style_run(run, bold=True, color=HEADER_GRAY, size=Pt(13))
    return p


def _add_label_value_table(doc, rows, label_width=Cm(4.5)):
    """rows: lista de tuplas (etiqueta, valor)."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table)
    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        row.cells[0].width = label_width
        _set_cell_text(row.cells[0], label, bold=True, shading=LABEL_GREEN)
        _set_cell_text(row.cells[1], value or "Por definir", shading=WHITE)
    doc.add_paragraph()
    return table


def _add_bullet_paragraphs(cell, items):
    cell.text = ""
    first = True
    for item in items:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        run = p.add_run(f"• {item}")
        _style_run(run)


# --------------------------------------------------------------------------
# Encabezado con logo
# --------------------------------------------------------------------------

def _add_header(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists(LOGO_PATH):
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.1))


def _sanitize_filename(text):
    text = re.sub(r"[^\w\s-]", "", text or "").strip()
    text = re.sub(r"[\s]+", "_", text)
    return text or "SinNombre"


# --------------------------------------------------------------------------
# Documento principal
# --------------------------------------------------------------------------

def generar_docx(perfil: dict, empresa: str, cargo: str, consultor: str = "") -> tuple[bytes, str]:
    """Construye el .docx del Perfil de Cargo.

    Devuelve (bytes_del_archivo, nombre_de_archivo_sugerido).
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE

    _add_header(doc)

    # Título del documento
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("PERFIL DE CARGO")
    _style_run(title_run, bold=True, size=Pt(16))
    title_p.paragraph_format.space_after = Pt(2)

    subtitle_p = doc.add_paragraph()
    subtitle_text = " / ".join([t for t in [empresa, cargo] if t])
    subtitle_run = subtitle_p.add_run(subtitle_text or "Cliente")
    _style_run(subtitle_run, bold=False, size=Pt(12), color=HEADER_GRAY)

    meta_p = doc.add_paragraph()
    meta_bits = [f"Fecha de levantamiento: {date.today().strftime('%d-%m-%Y')}"]
    if consultor:
        meta_bits.append(f"Consultor/a: {consultor}")
    meta_run = meta_p.add_run(" | ".join(meta_bits))
    _style_run(meta_run, size=Pt(9), color=HEADER_GRAY)
    meta_p.paragraph_format.space_after = Pt(10)

    # I. Datos Generales Empresa
    e = perfil.get("empresa", {})
    _add_heading(doc, "I", "DATOS GENERALES EMPRESA")
    _add_label_value_table(doc, [
        ("Definición Empresa", e.get("definicion_empresa")),
        ("Situación Actual", e.get("situacion_actual")),
        ("Área de la que Depende", e.get("area_de_la_que_depende")),
        ("Plazo Deseado de Ingreso", e.get("plazo_deseado_ingreso")),
        ("Opciones de Crecimiento", e.get("opciones_crecimiento")),
        ("Confidencialidad del Cargo", e.get("confidencialidad_cargo")),
    ])

    # II. Organigrama
    o = perfil.get("organigrama", {})
    _add_heading(doc, "II", "ORGANIGRAMA")
    _add_label_value_table(doc, [
        ("Nombre y Cargo de Jefatura Directa", o.get("jefatura_directa")),
        ("Reporta Indirectamente", o.get("reporta_indirectamente")),
        ("Personas a Cargo Directas", o.get("personas_a_cargo")),
        ("Tamaño de la Empresa", o.get("tamano_empresa")),
    ])

    # III. Descripción del Cargo
    dc = perfil.get("descripcion_cargo", {})
    _add_heading(doc, "III", "DESCRIPCIÓN DEL CARGO")
    table3 = doc.add_table(rows=4, cols=1)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table3)
    _set_cell_text(table3.rows[0].cells[0], "Propósito del Cargo", bold=True, shading=LABEL_GREEN)
    _set_cell_text(table3.rows[1].cells[0], dc.get("proposito_cargo") or "Por definir", shading=WHITE)
    _set_cell_text(table3.rows[2].cells[0], "Funciones del Cargo", bold=True, shading=LABEL_GREEN)
    funciones = dc.get("funciones_cargo") or []
    _shade_cell(table3.rows[3].cells[0], WHITE)
    if funciones:
        _add_bullet_paragraphs(table3.rows[3].cells[0], funciones)
    else:
        _set_cell_text(table3.rows[3].cells[0], "Por definir", shading=WHITE)
    doc.add_paragraph()

    # IV. Requisitos para el Cargo
    _add_heading(doc, "IV", "REQUISITOS PARA EL CARGO")
    requisitos = perfil.get("requisitos") or []
    req_by_name = {r.get("requerimiento", "").strip().lower(): r for r in requisitos}
    table4 = doc.add_table(rows=1 + len(REQUISITOS_FILAS), cols=3)
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table4)
    header_cells = table4.rows[0].cells
    for cell, text in zip(header_cells, ["Requerimiento", "Excluyente", "Deseable"]):
        _set_cell_text(cell, text, bold=True, shading=LABEL_GREEN)
    for i, nombre_fila in enumerate(REQUISITOS_FILAS, start=1):
        r = req_by_name.get(nombre_fila.strip().lower(), {})
        cells = table4.rows[i].cells
        _set_cell_text(cells[0], nombre_fila, bold=True, shading=LABEL_GREEN)
        _set_cell_text(cells[1], r.get("excluyente") or "No requerido", shading=WHITE)
        _set_cell_text(cells[2], r.get("deseable") or "No requerido", shading=WHITE)
    doc.add_paragraph()

    # V. Perfil Candidato
    _add_heading(doc, "V", "PERFIL CANDIDATO")
    _add_label_value_table(doc, [
        ("Características específicas que debe tener", perfil.get("perfil_candidato")),
    ])

    # VI. Competencias
    _add_heading(doc, "VI", "COMPETENCIAS")
    competencias = perfil.get("competencias") or []
    table6 = doc.add_table(rows=1 + max(len(competencias), 1), cols=2)
    table6.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table6)
    _set_cell_text(table6.rows[0].cells[0], "COMPETENCIA", bold=True, shading=LABEL_GREEN)
    _set_cell_text(table6.rows[0].cells[1], "DEFINICIÓN", bold=True, shading=LABEL_GREEN)
    if competencias:
        for i, c in enumerate(competencias, start=1):
            cells = table6.rows[i].cells
            _set_cell_text(cells[0], c.get("competencia") or "", bold=True, shading=WHITE)
            _set_cell_text(cells[1], c.get("definicion") or "", shading=WHITE)
    else:
        _set_cell_text(table6.rows[1].cells[0], "Por definir", bold=True, shading=WHITE)
        _set_cell_text(table6.rows[1].cells[1], "Por definir", shading=WHITE)
    doc.add_paragraph()

    # VII. Condiciones Laborales
    cl = perfil.get("condiciones_laborales", {})
    _add_heading(doc, "VII", "CONDICIONES LABORALES")
    _add_label_value_table(doc, [
        ("Ubicación", cl.get("ubicacion")),
        ("Jornada Laboral", cl.get("jornada_laboral")),
        ("Renta", cl.get("renta")),
        ("Beneficios de la Empresa", cl.get("beneficios")),
        ("Tipo de Contrato", cl.get("tipo_contrato")),
    ])

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"Perfil_Cargo_{_sanitize_filename(empresa)}_{_sanitize_filename(cargo)}_{date.today().strftime('%Y%m%d')}.docx"
    return buffer.read(), filename
